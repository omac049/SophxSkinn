"""Generate the SophxSkinn Brand Guide as a professional .docx document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

WORKSPACE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUTPUT = os.path.join(WORKSPACE, "SophxSkinn_Brand_Guide.docx")

# ── Brand colors ──
COLORS = {
    "Warm White":   {"hex": "#FFFBF9", "rgb": (255, 251, 249), "use": "Page backgrounds, light surfaces"},
    "Cream":        {"hex": "#F7F2EE", "rgb": (247, 242, 238), "use": "Primary background, body fill"},
    "Cream Dark":   {"hex": "#EDE6E0", "rgb": (237, 230, 224), "use": "Subtle dividers, muted surfaces"},
    "Blush 50":     {"hex": "#FDF6F3", "rgb": (253, 246, 243), "use": "Card backgrounds, soft highlights"},
    "Blush 100":    {"hex": "#F2D8D0", "rgb": (242, 216, 208), "use": "Accent tints, decorative elements"},
    "Blush 200":    {"hex": "#E0B8B0", "rgb": (224, 184, 176), "use": "Borders, soft accent lines"},
    "Rose 300":     {"hex": "#C49188", "rgb": (196, 145, 136), "use": "Secondary accent, gradient stops"},
    "Rose 500":     {"hex": "#A86A60", "rgb": (168, 106, 96),  "use": "Primary accent, eyebrow text, links"},
    "Rose 600":     {"hex": "#8F544A", "rgb": (143, 84, 74),   "use": "Primary CTA, headings, hover states"},
    "Berry 700":    {"hex": "#5C3432", "rgb": (92, 52, 50),    "use": "Signature text, deep emphasis"},
    "Ink 500":      {"hex": "#7A6864", "rgb": (122, 104, 100), "use": "Secondary body text, captions"},
    "Ink 700":      {"hex": "#4D3D39", "rgb": (77, 61, 57),    "use": "Navigation links, subtle headings"},
    "Ink 800":      {"hex": "#2C221F", "rgb": (44, 34, 31),    "use": "Dark gradients, footer background"},
    "Ink 900":      {"hex": "#181210", "rgb": (24, 18, 16),    "use": "Primary body text, hero overlay, footer"},
}

INK_900 = RGBColor(24, 18, 16)
ROSE_500 = RGBColor(168, 106, 96)
ROSE_600 = RGBColor(143, 84, 74)
INK_500 = RGBColor(122, 104, 100)
CREAM = RGBColor(247, 242, 238)
WHITE = RGBColor(255, 255, 255)
BERRY = RGBColor(92, 52, 50)

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    hex_color = hex_color.lstrip("#")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, top=None, bottom=None, start=None, end=None):
    """Set cell borders. Each kwarg is a dict with sz, color, val keys."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for side, props in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        if props:
            el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="{props.get("val","single")}" '
                f'w:sz="{props.get("sz","4")}" w:space="0" '
                f'w:color="{props.get("color","auto")}"/>'
            )
            borders.append(el)
    tcPr.append(borders)

def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)

def add_styled_paragraph(doc, text, font_name="Calibri", size=11, bold=False,
                         color=INK_900, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=0, space_after=6, italic=False,
                         all_caps=False, letter_spacing=None):
    """Add a paragraph with specific styling."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if all_caps:
        run.font.all_caps = True
    if letter_spacing:
        run.font.cs_bold = False
        rpr = run._element.get_or_add_rPr()
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:val="{letter_spacing}"/>')
        rpr.append(spacing)
    return p

def add_section_title(doc, title):
    """Add a major section title with consistent styling."""
    add_styled_paragraph(doc, "", size=6, space_after=0)
    p = add_styled_paragraph(
        doc, title.upper(),
        font_name="Calibri", size=22, bold=True,
        color=INK_900, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=24, space_after=4,
        all_caps=True, letter_spacing=120
    )
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(12)
    run = sep.add_run()
    # thin rose line
    from docx.oxml import OxmlElement
    pPr = sep._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A86A60")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_subsection_title(doc, title):
    """Add a subsection title."""
    add_styled_paragraph(
        doc, title,
        font_name="Calibri", size=14, bold=True,
        color=ROSE_600, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=16, space_after=6
    )

def add_body_text(doc, text, space_after=8):
    """Add body text."""
    return add_styled_paragraph(
        doc, text,
        font_name="Calibri", size=10.5, bold=False,
        color=INK_900, space_after=space_after
    )

def add_quote(doc, text):
    """Add a styled quote/tagline."""
    add_styled_paragraph(
        doc, f'"{text}"',
        font_name="Georgia", size=14, bold=False, italic=True,
        color=ROSE_600, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=12
    )

def add_page_break(doc):
    doc.add_page_break()

def build_cover_page(doc):
    """Build the title/cover page."""
    for _ in range(6):
        add_styled_paragraph(doc, "", size=12, space_after=0)

    add_styled_paragraph(
        doc, "SOPHXSKINN",
        font_name="Calibri", size=42, bold=True,
        color=INK_900, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=4,
        letter_spacing=200
    )

    add_styled_paragraph(
        doc, "Brand Guide",
        font_name="Georgia", size=20, bold=False, italic=True,
        color=ROSE_500, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=24
    )

    # thin separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(24)
    run = sep.add_run("_" * 40)
    run.font.size = Pt(1)
    run.font.color.rgb = ROSE_500

    add_styled_paragraph(
        doc, "Skin rituals with a fashion-forward edge.",
        font_name="Georgia", size=12, italic=True,
        color=INK_500, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=6
    )

    add_styled_paragraph(
        doc, "Facials  |  Brows  |  Lashes  |  Waxing",
        font_name="Calibri", size=10, bold=False,
        color=INK_500, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=48
    )

    add_styled_paragraph(
        doc, "Licensed Esthetician",
        font_name="Calibri", size=10, bold=True,
        color=ROSE_600, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=2,
        all_caps=True, letter_spacing=100
    )

    add_styled_paragraph(
        doc, "The Remedy Salon  \u00b7  Goodyear, Arizona",
        font_name="Calibri", size=10,
        color=INK_500, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=2
    )

    add_styled_paragraph(
        doc, "sophxskinn.com",
        font_name="Calibri", size=10,
        color=ROSE_500, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=0
    )


def build_toc(doc):
    """Build a table of contents page."""
    add_page_break(doc)

    add_styled_paragraph(
        doc, "CONTENTS",
        font_name="Calibri", size=18, bold=True,
        color=INK_900, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=24, space_after=20,
        all_caps=True, letter_spacing=120
    )

    sections = [
        ("01", "Brand Overview"),
        ("02", "Brand Pillars & Positioning"),
        ("03", "Logo & Wordmark"),
        ("04", "Color Palette"),
        ("05", "Typography"),
        ("06", "Voice & Tone"),
        ("07", "Photography & Visual Style"),
        ("08", "UI & Digital Patterns"),
        ("09", "Social Media Guidelines"),
        ("10", "Contact & Assets"),
    ]

    for num, title in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        run_num = p.add_run(f"{num}   ")
        run_num.font.name = "Calibri"
        run_num.font.size = Pt(10)
        run_num.font.bold = True
        run_num.font.color.rgb = ROSE_500
        run_title = p.add_run(title)
        run_title.font.name = "Calibri"
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = INK_900


def build_brand_overview(doc):
    """Section 1: Brand Overview."""
    add_page_break(doc)
    add_section_title(doc, "01  Brand Overview")

    add_subsection_title(doc, "Brand Name")
    add_body_text(doc, 'The brand name is "SophxSkinn" -- always written with a capital S, lowercase oph, capital X (no space), capital S, lowercase kinn. The "x" acts as a stylistic connector between the personal identity and the skincare discipline.')

    add_subsection_title(doc, "Descriptor")
    add_body_text(doc, "Licensed Esthetician")

    add_subsection_title(doc, "Location")
    add_body_text(doc, "The Remedy Salon")
    add_body_text(doc, "13375 W McDowell Rd #108, Goodyear, AZ 85395")

    add_subsection_title(doc, "Primary Tagline")
    add_quote(doc, "Skin rituals with a fashion-forward edge.")

    add_subsection_title(doc, "Editorial Tagline")
    add_quote(doc, "Where skincare gets its own runway moment.")

    add_subsection_title(doc, "Hero Headline")
    add_quote(doc, "Beauty with intention. Edge included.")

    add_subsection_title(doc, "Services Offered")
    services = ["Facials -- Hydration, texture, and the kind of glow that turns heads.",
                "Brows -- Sculpted arches that frame your face like the accessory it deserves.",
                "Lashes -- The finishing touch your look has been waiting for -- effortless lift, serious definition.",
                "Waxing -- Polished, smooth skin with a comfort-first touch and luxe aftercare."]
    for s in services:
        add_body_text(doc, f"  \u2022  {s}")

    add_subsection_title(doc, "Website")
    add_body_text(doc, "sophxskinn.com")

    add_subsection_title(doc, "Booking Platform")
    add_body_text(doc, "Vagaro (vagaro.com/theremedy)")


def build_brand_pillars(doc):
    """Section 2: Brand Pillars & Positioning."""
    add_page_break(doc)
    add_section_title(doc, "02  Brand Pillars & Positioning")

    add_body_text(doc,
        "SophxSkinn is built on three foundational pillars that guide every client interaction, "
        "visual decision, and piece of communication. These pillars should be evident in all brand touchpoints."
    )

    pillars = [
        ("Precision", "Every service is executed with meticulous care and technical skill. From brow arches to facial protocols, nothing is rushed or generic. Precision is the standard."),
        ("Polish", "The brand exudes a refined, fashion-forward sensibility. Polish shows up in the visual identity, the studio environment, the client experience, and the finished results. Every detail is considered."),
        ("Intention", "Nothing about SophxSkinn is accidental. Each appointment is tailored, each product is chosen with purpose, and each client relationship is built with genuine care. Intention is what transforms a service into a ritual."),
    ]

    for title, desc in pillars:
        add_subsection_title(doc, title)
        add_body_text(doc, desc)

    add_subsection_title(doc, "Brand Positioning Statement")
    add_body_text(doc,
        "SophxSkinn is a boutique esthetician brand where skincare and fashion sensibility are "
        "inseparable. Based at The Remedy Salon in Goodyear, Arizona, SophxSkinn delivers facials, "
        "brows, lashes, and waxing with the same intentionality you'd expect from a personal stylist -- "
        "precise technique, polished results, and a playful confidence that never takes itself too seriously."
    )

    add_subsection_title(doc, "Target Audience")
    add_body_text(doc,
        "Style-obsessed individuals who see skincare as part of getting dressed -- not separate from it. "
        "They follow beauty editors and fashion accounts equally. They want an esthetician who gets "
        "both the science and the aesthetic, and they'd rather book with someone who has a point of view "
        "than a generic spa. They value personality alongside professionalism."
    )


def build_logo_section(doc):
    """Section 3: Logo & Wordmark."""
    add_page_break(doc)
    add_section_title(doc, "03  Logo & Wordmark")

    add_subsection_title(doc, "Primary Logo (Script Wordmark)")
    add_body_text(doc,
        'The primary SophxSkinn logo is a script-style wordmark (file: Sophxskinn.png). '
        'It is used as the hero brand mark on the website, in the site header, the footer, '
        'and as the Open Graph / social sharing image.'
    )
    add_body_text(doc,
        "The script logo conveys warmth, personal artistry, and approachability while "
        "maintaining a polished, editorial feel consistent with the brand's fashion-forward positioning."
    )

    add_subsection_title(doc, "Secondary Mark / Favicon")
    add_body_text(doc,
        "A secondary graphic mark is used as the browser favicon and as an alternate "
        "Open Graph image on interior pages (file: AB8A5850-BC05-46A2-90A2-BAFFF757C5E9.png). "
        "This mark works at small sizes and serves as a compact brand identifier."
    )

    add_subsection_title(doc, "Logo Color Treatments")

    treatments = [
        ("On light backgrounds (cream/white)", "Display the logo in its original colorway (warm tones). The logo should contrast clearly against Cream (#F7F2EE) or Warm White (#FFFBF9) backgrounds."),
        ("On dark backgrounds (hero, footer)", "Apply a white treatment using CSS filter: brightness(10) saturate(0). The logo reads as a clean white mark on dark imagery or the Ink 900 (#181210) footer."),
        ("On photography", "Use the white treatment over darkened photography with a gradient veil overlay to ensure legibility."),
    ]
    for label, desc in treatments:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run_label = p.add_run(f"{label}: ")
        run_label.font.name = "Calibri"
        run_label.font.size = Pt(10.5)
        run_label.font.bold = True
        run_label.font.color.rgb = INK_900
        run_desc = p.add_run(desc)
        run_desc.font.name = "Calibri"
        run_desc.font.size = Pt(10.5)
        run_desc.font.color.rgb = INK_900

    add_subsection_title(doc, "Clear Space & Minimum Size")
    add_body_text(doc,
        "Maintain a clear space around the logo equal to the height of the 'x' in the wordmark "
        "on all sides. The logo should never appear smaller than 120px wide on digital screens "
        "or 1.5 inches in print to preserve legibility of the script letterforms."
    )

    add_subsection_title(doc, "Logo Misuse")
    misuses = [
        "Do not stretch, skew, or distort the logo proportions.",
        "Do not place the logo on busy or low-contrast backgrounds without a veil or container.",
        "Do not alter the logo colors outside the approved treatments above.",
        "Do not add drop shadows, outlines, or effects to the logo.",
        "Do not recreate the logo in a substitute typeface.",
        "Do not rotate the logo.",
    ]
    for m in misuses:
        add_body_text(doc, f"  \u2022  {m}", space_after=4)


def build_color_palette(doc):
    """Section 4: Color Palette."""
    add_page_break(doc)
    add_section_title(doc, "04  Color Palette")

    add_body_text(doc,
        "The SophxSkinn color palette is rooted in warm, skin-inspired tones. "
        "It pairs soft creams and blushes with deep, grounding inks and a signature dusty rose accent. "
        "The palette evokes warmth, sophistication, and tactile comfort."
    )

    # Primary palette
    add_subsection_title(doc, "Primary Palette")
    add_body_text(doc, "These are the core colors used across all brand materials.")

    primary_colors = ["Rose 600", "Rose 500", "Rose 300", "Ink 900", "Cream", "Warm White"]
    table = doc.add_table(rows=len(primary_colors) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    headers = ["Color", "Hex", "RGB", "Usage"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = INK_500
        run.font.all_caps = True
        set_cell_border(cell, bottom={"sz": "4", "color": "A86A60"})

    for idx, name in enumerate(primary_colors, 1):
        c = COLORS[name]
        row = table.rows[idx]
        values = [name, c["hex"], f"({c['rgb'][0]}, {c['rgb'][1]}, {c['rgb'][2]})", c["use"]]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            run.font.color.rgb = INK_900
            if i == 0:
                run.font.bold = True
            set_cell_border(cell, bottom={"sz": "2", "color": "EDE6E0"})

        swatch_cell = row.cells[0]
        set_cell_shading(swatch_cell, c["hex"])
        if name in ["Ink 900", "Berry 700"]:
            for p in swatch_cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE

    # Extended palette
    add_styled_paragraph(doc, "", size=6, space_after=0)
    add_subsection_title(doc, "Extended Palette")
    add_body_text(doc, "Supporting colors for subtle UI elements, borders, and secondary surfaces.")

    extended_colors = ["Blush 50", "Blush 100", "Blush 200", "Cream Dark", "Berry 700", "Ink 500", "Ink 700", "Ink 800"]
    table2 = doc.add_table(rows=len(extended_colors) + 1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table2)

    for i, h in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = INK_500
        run.font.all_caps = True
        set_cell_border(cell, bottom={"sz": "4", "color": "A86A60"})

    for idx, name in enumerate(extended_colors, 1):
        c = COLORS[name]
        row = table2.rows[idx]
        values = [name, c["hex"], f"({c['rgb'][0]}, {c['rgb'][1]}, {c['rgb'][2]})", c["use"]]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            run.font.color.rgb = INK_900
            if i == 0:
                run.font.bold = True
            set_cell_border(cell, bottom={"sz": "2", "color": "EDE6E0"})

        swatch_cell = row.cells[0]
        set_cell_shading(swatch_cell, c["hex"])
        if name in ["Ink 800", "Ink 700", "Berry 700", "Ink 500"]:
            for p in swatch_cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE

    add_styled_paragraph(doc, "", size=6, space_after=0)
    add_subsection_title(doc, "Gradient Treatments")
    gradients = [
        ("Primary CTA Gradient", "135deg: Rose 600 (#8F544A) -> Rose 500 (#A86A60) -> Rose 300 (#C49188)", "Book buttons, primary action elements"),
        ("Scroll Progress Bar", "90deg: Rose 600 (#8F544A) -> Rose 300 (#C49188) -> Blush 100 (#F2D8D0)", "Scroll indicator at top of viewport"),
        ("Dark CTA / Support Strip", "140deg: Ink 900 (#181210) -> Ink 800 (#2C221F) -> Rose 600 (#8F544A)", "Dark promotional blocks, footer accents"),
        ("Card Surface", "168deg: Warm White at 98% -> Blush 100 at 32%", "Card backgrounds with subtle warmth"),
        ("Hero Veil", "180deg: Ink transparent -> Ink at 72%", "Darkening overlay on hero photography"),
    ]
    for name, spec, use in gradients:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run_name = p.add_run(f"{name}:  ")
        run_name.font.name = "Calibri"
        run_name.font.size = Pt(10)
        run_name.font.bold = True
        run_name.font.color.rgb = ROSE_600
        run_spec = p.add_run(f"{spec}  ")
        run_spec.font.name = "Calibri"
        run_spec.font.size = Pt(10)
        run_spec.font.color.rgb = INK_900
        run_use = p.add_run(f"({use})")
        run_use.font.name = "Calibri"
        run_use.font.size = Pt(9.5)
        run_use.font.italic = True
        run_use.font.color.rgb = INK_500


def build_typography(doc):
    """Section 5: Typography."""
    add_page_break(doc)
    add_section_title(doc, "05  Typography")

    add_body_text(doc,
        "SophxSkinn uses a curated type system pairing geometric sans-serif body text with "
        "editorial serifs for display and accent moments. The hierarchy creates a feeling of "
        "polish and editorial sophistication while maintaining excellent readability."
    )
    add_body_text(doc, "All typefaces are sourced from Google Fonts for consistent cross-platform rendering.")

    # Type table
    fonts = [
        ("Manrope", "Body Text", "400 - 700", "Primary body copy, descriptions, paragraphs, captions, UI text", "sans-serif"),
        ("Inter", "Headings & UI", "500 - 800", "H1 headlines, H3 subheadings, navigation links, bold UI labels, buttons", "sans-serif"),
        ("Playfair Display", "Display & Editorial", "400 - 700 (+ italic)", "H2 section headlines, editorial quotes, marquee text, large decorative type", "serif"),
        ("MonteCarlo", "Script Accent", "400", "Default H2 styling (when not overridden), decorative moments requiring warmth", "cursive"),
    ]

    table = doc.add_table(rows=len(fonts) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    t_headers = ["Typeface", "Role", "Weights", "Usage", "Fallback"]
    for i, h in enumerate(t_headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = INK_500
        run.font.all_caps = True
        set_cell_border(cell, bottom={"sz": "4", "color": "A86A60"})

    for idx, (name, role, weights, usage, fallback) in enumerate(fonts, 1):
        row = table.rows[idx]
        vals = [name, role, weights, usage, fallback]
        for i, val in enumerate(vals):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            run.font.color.rgb = INK_900
            if i == 0:
                run.font.bold = True
            set_cell_border(cell, bottom={"sz": "2", "color": "EDE6E0"})

    add_styled_paragraph(doc, "", size=6, space_after=0)
    add_subsection_title(doc, "Type Scale & Hierarchy")

    scales = [
        ("H1 (Hero Headline)", "Inter 800, clamp(2.8rem - 5.2rem), line-height 0.98, tracking -0.035em", "Used once per page for the primary statement"),
        ("H2 (Section Headline)", "Playfair Display 600, clamp(2.6rem - 4.8rem), line-height 0.88, tracking -0.02em, Rose 600 color", "Major section introductions; slight -0.3deg rotation for editorial flair"),
        ("H2 Default (Script)", "MonteCarlo 400, line-height 1.02", "Used when Playfair Display is not explicitly applied"),
        ("H3 (Subsection)", "Inter 600, line-height 1.08, tracking -0.01em", "Card titles, feature names, FAQ questions"),
        ("Eyebrow", "0.72rem, uppercase, 700 weight, tracking 0.2em, Rose 500 color", "Section labels above headlines; context cues"),
        ("Body", "Manrope 400-500, 1rem base, line-height 1.6-1.65", "Paragraphs, descriptions, general content"),
        ("Caption / Small", "0.84-0.86rem, 600 weight, tracking 0.03-0.04em", "Image captions, fine print, pillar text"),
        ("Button Text", "0.92rem, 600 weight, tracking 0.02em", "All interactive CTAs and navigation links"),
    ]

    for label, spec, note in scales:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run_label = p.add_run(f"{label}:  ")
        run_label.font.name = "Calibri"
        run_label.font.size = Pt(10)
        run_label.font.bold = True
        run_label.font.color.rgb = INK_900
        run_spec = p.add_run(spec)
        run_spec.font.name = "Calibri"
        run_spec.font.size = Pt(10)
        run_spec.font.color.rgb = INK_500
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
        p2.paragraph_format.left_indent = Inches(0.3)
        run_note = p2.add_run(note)
        run_note.font.name = "Calibri"
        run_note.font.size = Pt(9.5)
        run_note.font.italic = True
        run_note.font.color.rgb = INK_500


def build_voice_tone(doc):
    """Section 6: Voice & Tone."""
    add_page_break(doc)
    add_section_title(doc, "06  Voice & Tone")

    add_body_text(doc,
        "The SophxSkinn voice is the intersection of a fashion editor and your coolest friend -- "
        "polished enough for a luxury brand, playful enough to feel like a real person is behind it. "
        "The tone has a wink without trying too hard. Copy should feel like scrolling a curated fashion "
        "feed, not reading a spa brochure. High-end vocabulary, low-effort cool."
    )

    add_subsection_title(doc, "Voice Attributes")
    attributes = [
        ("Confident with a wink", "We know what we're good at and we'll say it -- with just enough playfulness to keep it charming, never arrogant."),
        ("Polished, not stiff", "The tone is elevated and refined, but it breathes. Sentences can be short. Fragments are fine. Personality is non-negotiable."),
        ("Fashion-fluent", "We borrow from the language of style -- 'edit', 'moment', 'finishing touch', 'accessory'. Skincare is fashion's next chapter."),
        ("Playful, not juvenile", "A subtle wink, not a loud laugh. Think Glossier's warmth meets Jacquemus' confidence. Self-aware, never try-hard."),
        ("Personal, not generic", "We speak as a named esthetician with a point of view and actual taste -- not a faceless corporation with stock-photo energy."),
    ]
    for title, desc in attributes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run_t = p.add_run(f"{title}:  ")
        run_t.font.name = "Calibri"
        run_t.font.size = Pt(10.5)
        run_t.font.bold = True
        run_t.font.color.rgb = ROSE_600
        run_d = p.add_run(desc)
        run_d.font.name = "Calibri"
        run_d.font.size = Pt(10.5)
        run_d.font.color.rgb = INK_900

    add_subsection_title(doc, "Key Phrases & Language")
    phrases = [
        '"Skin rituals" (not "skin treatments" or "procedures")',
        '"Fashion-forward edge" (not "beautiful results")',
        '"Glow edit" (not "service menu")',
        '"Runway moment" / "style upgrade" (not "makeover")',
        '"Soft-glam edge" (not "natural look")',
        '"The finishing touch" (not "add-on service")',
        '"Ritual" over "appointment", "moment" over "experience"',
        '"Edge included" -- the brand\'s signature throwaway confidence',
        '"Obsession-worthy" / "turns heads" -- aspirational but grounded',
    ]
    for ph in phrases:
        add_body_text(doc, f"  \u2022  {ph}", space_after=4)

    add_subsection_title(doc, "Sample Copy Patterns")
    samples = [
        ("Service description", '"Hydration, texture, and the kind of glow that turns heads."'),
        ("Process step", '"Settle into The Remedy Salon for a calm, fashion-forward treatment tailored to you."'),
        ("CTA context", '"Live openings, easy checkout, instant confirmation -- booking should feel as good as the appointment."'),
        ("Playful aside", '"Most clients book in under 60 seconds. Seriously."'),
        ("Micro-copy / caption", '"The vibe? Polished." / "Edge included."'),
    ]
    for label, sample in samples:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run_l = p.add_run(f"{label}:  ")
        run_l.font.name = "Calibri"
        run_l.font.size = Pt(10)
        run_l.font.bold = True
        run_l.font.color.rgb = INK_900
        run_s = p.add_run(sample)
        run_s.font.name = "Georgia"
        run_s.font.size = Pt(10)
        run_s.font.italic = True
        run_s.font.color.rgb = ROSE_600


def build_photography(doc):
    """Section 7: Photography & Visual Style."""
    add_page_break(doc)
    add_section_title(doc, "07  Photography & Visual Style")

    add_subsection_title(doc, "Photography Direction")
    add_body_text(doc,
        "SophxSkinn photography is warm, editorial, and intentional. Images should feel like "
        "pages from a beauty-meets-fashion editorial -- never stock, never clinical. "
        "Natural light is preferred, with warm tones that complement the brand palette."
    )

    add_subsection_title(doc, "Image Categories")
    categories = [
        ("Esthetician Portraits", "Professional yet approachable portraits of the esthetician. These images anchor the personal brand. Warm lighting, soft backgrounds, styled but not overly posed. Used in hero sections, meet sections, and about content."),
        ("Client Results", "Close-up and detail shots showcasing brow work, lash lifts, facial glow, and skin texture. Clean framing with consistent warm-tone editing. Used in service tiles, mosaics, and social proof."),
        ("Studio / Session Videos", "Short-form MP4 video clips of real sessions (facials, brow shaping, lash treatments). These play as ambient content in image mosaics with auto-play, muted, and looped."),
    ]
    for title, desc in categories:
        add_styled_paragraph(
            doc, title,
            font_name="Calibri", size=11, bold=True,
            color=ROSE_600, space_before=8, space_after=4
        )
        add_body_text(doc, desc)

    add_subsection_title(doc, "Visual Treatments")
    treatments = [
        ("Gradient overlays (veils)", "Dark gradient scrims over photography for text legibility. Primary veil: linear-gradient from transparent ink to 72% ink, bottom-heavy."),
        ("Rounded containers", "All image tiles use large border-radius: 28px (XL), 20px (LG), or 12px (MD). Squared or sharp-cornered images are not on-brand."),
        ("Hover zoom", "Image tiles scale subtly on hover (1.04x) for interactivity without jarring effects."),
        ("Shadow layers", "Soft, warm-toned shadows using rgba(24, 18, 16, ...) at low opacity. Never harsh or cool-toned drop shadows."),
        ("Blush orbs", "Decorative radial gradient circles (rgba(222, 181, 173, 0.3)) that float with a slow animation, adding depth and softness to sections."),
        ("Backdrop blur / Glass effect", "Header and select buttons use backdrop-filter: blur(24px) saturate(1.6) for a frosted-glass surface treatment."),
    ]
    for title, desc in treatments:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run_t = p.add_run(f"{title}:  ")
        run_t.font.name = "Calibri"
        run_t.font.size = Pt(10.5)
        run_t.font.bold = True
        run_t.font.color.rgb = INK_900
        run_d = p.add_run(desc)
        run_d.font.name = "Calibri"
        run_d.font.size = Pt(10.5)
        run_d.font.color.rgb = INK_900

    add_subsection_title(doc, "Photo Editing Guidelines")
    edits = [
        "Maintain warm white balance consistent with the cream/blush palette.",
        "Avoid heavy filters, blue/cool shifts, or over-saturated skin tones.",
        "Skin retouching should be minimal and natural -- soft-glam, never airbrushed.",
        "Background tones should complement Cream (#F7F2EE) or Ink 900 (#181210).",
        "All images should feel cohesive when placed side by side in a grid or mosaic.",
    ]
    for e in edits:
        add_body_text(doc, f"  \u2022  {e}", space_after=4)


def build_ui_patterns(doc):
    """Section 8: UI & Digital Patterns."""
    add_page_break(doc)
    add_section_title(doc, "08  UI & Digital Patterns")

    add_subsection_title(doc, "Layout System")
    layout_specs = [
        ("Max content width", "1200px"),
        ("Side margins", "24px per side (48px total gap from edges)"),
        ("Section spacing", "clamp(2.5rem, 4.5vw, 4.5rem) between major sections"),
        ("Card padding", "clamp(2.2rem, 4vw, 3.5rem)"),
    ]
    for label, val in layout_specs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        run_l = p.add_run(f"{label}:  ")
        run_l.font.name = "Calibri"
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        run_l.font.color.rgb = INK_900
        run_v = p.add_run(val)
        run_v.font.name = "Calibri"
        run_v.font.size = Pt(10.5)
        run_v.font.color.rgb = INK_500

    add_subsection_title(doc, "Border Radius System")
    radii = [
        ("XL (28px)", "Cards, content sections, image containers, hero elements"),
        ("LG (20px)", "Mosaic tiles, medium containers, photo frames"),
        ("MD (12px)", "Support strip items, small UI cards, inline containers"),
        ("Pill (999px)", "Buttons, pills, tags, skip link, navigation toggles"),
    ]
    for label, use in radii:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        run_l = p.add_run(f"{label}:  ")
        run_l.font.name = "Calibri"
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        run_l.font.color.rgb = INK_900
        run_u = p.add_run(use)
        run_u.font.name = "Calibri"
        run_u.font.size = Pt(10.5)
        run_u.font.color.rgb = INK_500

    add_subsection_title(doc, "Shadow System")
    shadows = [
        ("Soft", "0 4px 24px rgba(24,18,16, 0.055) -- default card resting state"),
        ("Medium", "0 10px 36px rgba(24,18,16, 0.075) -- mosaic tiles, moderate emphasis"),
        ("Strong", "0 20px 52px rgba(24,18,16, 0.11) -- hero images, featured elements"),
        ("Card Hover", "0 10px 34px rgba(24,18,16, 0.095) -- interactive lift on hover"),
    ]
    for label, spec in shadows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        run_l = p.add_run(f"{label}:  ")
        run_l.font.name = "Calibri"
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        run_l.font.color.rgb = INK_900
        run_s = p.add_run(spec)
        run_s.font.name = "Calibri"
        run_s.font.size = Pt(10.5)
        run_s.font.color.rgb = INK_500

    add_subsection_title(doc, "Button Styles")

    add_styled_paragraph(doc, "Primary (Book Button)", font_name="Calibri", size=11, bold=True, color=ROSE_600, space_before=8, space_after=4)
    add_body_text(doc, "Pill shape (999px radius), white text on rose gradient background (Rose 600 -> Rose 500 -> Rose 300 at 135deg). 1px Rose 600 border. Box shadow with warm rose-tinted glow. Hover: lifts 2px with deeper shadow. Shimmer sweep animation on hover.")

    add_styled_paragraph(doc, "Ghost Button", font_name="Calibri", size=11, bold=True, color=ROSE_600, space_before=8, space_after=4)
    add_body_text(doc, "Pill shape, Rose 600 text on semi-transparent warm white background (rgba 255,250,248, 0.88). 1px Rose 300 border. Hover: lifts 2px, fills solid warm white, border shifts to Rose 500.")

    add_styled_paragraph(doc, "Hero Variant (Primary)", font_name="Calibri", size=11, bold=True, color=ROSE_600, space_before=8, space_after=4)
    add_body_text(doc, "Inverted for dark hero backgrounds: Warm White background with Rose 600 text. No visible border. Deep warm shadow.")

    add_styled_paragraph(doc, "Hero Variant (Ghost)", font_name="Calibri", size=11, bold=True, color=ROSE_600, space_before=8, space_after=4)
    add_body_text(doc, "Frosted glass: semi-transparent white background (8% opacity) with backdrop-filter blur. White text. Subtle white border at 30% opacity.")

    add_subsection_title(doc, "Header / Navigation")
    add_body_text(doc, "Fixed position, 72px height. Default state: frosted cream background (rgba 247,242,239, 0.82) with backdrop-filter blur(24px) saturate(1.6). Transparent variant over hero sections with white nav text. Navigation links: uppercase, 600 weight, 0.04em tracking, 0.92rem. Active/hover: Rose 600 color with animated underline.")

    add_subsection_title(doc, "Scroll Progress Bar")
    add_body_text(doc, "Fixed 2px bar at viewport top. Rose gradient (Rose 600 -> Rose 300 -> Blush 100) that scales horizontally with scroll position. z-index 220.")

    add_subsection_title(doc, "Motion & Animation")
    animations = [
        "Page transitions via Framer Motion with staggered children (0.12s stagger, 0.6s container).",
        "Entry animations: fade-in with 24px upward translate, 0.7s duration, cubic-bezier(0.22, 1, 0.36, 1) easing.",
        "Hero image: subtle scale-down from 1.08 to 1.0 on load (1.4s duration).",
        "Button hover: translateY(-2px) with scale(1.03), 0.3s transition.",
        "Floating blush orbs: 12-14s ease-in-out infinite keyframe animation.",
        "All motion respects prefers-reduced-motion: falls back to simple opacity fades.",
    ]
    for a in animations:
        add_body_text(doc, f"  \u2022  {a}", space_after=4)


def build_social_media(doc):
    """Section 9: Social Media."""
    add_page_break(doc)
    add_section_title(doc, "09  Social Media Guidelines")

    add_subsection_title(doc, "Primary Platform")
    add_body_text(doc, "Instagram (@sophxskinn) -- instagram.com/sophxskinn/")

    add_subsection_title(doc, "Handle & Naming")
    add_body_text(doc, 'Use "@sophxskinn" consistently across all platforms. Do not abbreviate or alter the handle. The display name should read "SophxSkinn" with proper capitalization.')

    add_subsection_title(doc, "Content Pillars for Social")
    content_pillars = [
        ("Before & After / The Receipts", "Client transformations that let the work speak for itself. Warm editing, clean framing, minimal text overlays. The results are the caption."),
        ("Studio Diaries", "Behind-the-scenes glimpses of sessions, products, and the ritual in motion. Shows the craft without over-explaining. Reinforces 'intention' and 'precision'."),
        ("Skin Meets Style", "Where aesthetics content crosses into fashion territory -- outfit-of-the-day energy applied to skincare, GRWM-adjacent, editorial mood. The signature SophxSkinn lane."),
        ("The Person Behind the Brand", "Personality-forward content that shows taste, humor, and real life. Fashion pulls, studio styling, unfiltered moments. Polished, never performative."),
    ]
    for title, desc in content_pillars:
        add_styled_paragraph(doc, title, font_name="Calibri", size=11, bold=True, color=ROSE_600, space_before=6, space_after=3)
        add_body_text(doc, desc)

    add_subsection_title(doc, "Hashtag Strategy")
    add_body_text(doc, "Primary: #SophxSkinn #TheRemedySalon #GoodyearEsthetician")
    add_body_text(doc, "Service-specific: #BrowLamination #LashLift #FacialGlow #WaxingGoodyear")
    add_body_text(doc, "Community: #AZEsthetician #GoodyearAZ #ArizonaBeauty #SkincareRoutine")

    add_subsection_title(doc, "Visual Consistency")
    add_body_text(doc,
        "Social posts should maintain the brand's warm, editorial aesthetic. Use consistent photo editing "
        "that aligns with the cream/blush/rose palette. Grid layout should feel curated and intentional, "
        "alternating between close-up details, portraits, and text-based content."
    )


def build_contact_assets(doc):
    """Section 10: Contact & Assets."""
    add_page_break(doc)
    add_section_title(doc, "10  Contact & Assets")

    add_subsection_title(doc, "Digital Presence")
    contacts = [
        ("Website", "sophxskinn.com"),
        ("Instagram", "@sophxskinn (instagram.com/sophxskinn/)"),
        ("Booking", "vagaro.com/theremedy"),
        ("Location", "The Remedy Salon, 13375 W McDowell Rd #108, Goodyear, AZ 85395"),
    ]
    for label, val in contacts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        run_l = p.add_run(f"{label}:  ")
        run_l.font.name = "Calibri"
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        run_l.font.color.rgb = INK_900
        run_v = p.add_run(val)
        run_v.font.name = "Calibri"
        run_v.font.size = Pt(10.5)
        run_v.font.color.rgb = ROSE_500

    add_subsection_title(doc, "Brand Asset Files")
    add_body_text(doc, "The following files constitute the core brand asset library:")

    assets = [
        ("Sophxskinn.png", "Primary script logo / wordmark. Used in header, hero, footer, OG images, and structured data."),
        ("Favicon mark", "Compact brand identifier for browser tabs and secondary OG images."),
        ("Esthetician portraits (7)", "Professional photo series (sophxskinn-1.jpg through sophxskinn-7.jpg). Anchor the personal brand across hero, gallery, and meet sections."),
        ("Client result photos (5)", "Service result documentation: brow, brow lamination, lash, lash lift, and facial glow images."),
        ("Session videos (3)", "MP4 clips: facial session, brow shaping, lash treatment. Used in ambient mosaic displays."),
    ]
    for name, desc in assets:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run_n = p.add_run(f"{name}:  ")
        run_n.font.name = "Calibri"
        run_n.font.size = Pt(10.5)
        run_n.font.bold = True
        run_n.font.color.rgb = INK_900
        run_d = p.add_run(desc)
        run_d.font.name = "Calibri"
        run_d.font.size = Pt(10.5)
        run_d.font.color.rgb = INK_900

    add_subsection_title(doc, "SEO & Technical Brand Data")
    seo_items = [
        ("Theme color", "#1A1A1A (browser UI meta)"),
        ("Canonical domain", "sophxskinn.com"),
        ("Structured data", "JSON-LD LocalBusiness schema with logo, geo coordinates, and service descriptions"),
        ("OG image", "Sophxskinn.png script logo (home); favicon mark (services page)"),
        ("Meta description pattern", '"SophxSkinn -- Licensed esthetician at The Remedy Salon, [address]. [Services]. Book online."'),
    ]
    for label, val in seo_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        run_l = p.add_run(f"{label}:  ")
        run_l.font.name = "Calibri"
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        run_l.font.color.rgb = INK_900
        run_v = p.add_run(val)
        run_v.font.name = "Calibri"
        run_v.font.size = Pt(10.5)
        run_v.font.color.rgb = INK_500

    # Final footer
    add_styled_paragraph(doc, "", size=12, space_after=0)
    add_styled_paragraph(doc, "", size=12, space_after=0)
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from docx.oxml import OxmlElement
    pPr = sep._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C49188")
    pBdr.append(bottom)
    pPr.append(pBdr)

    add_styled_paragraph(
        doc, "SophxSkinn Brand Guide  |  Confidential",
        font_name="Calibri", size=9,
        color=INK_500, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=2
    )
    add_styled_paragraph(
        doc, "Prepared for internal use and approved partners.",
        font_name="Calibri", size=9,
        color=INK_500, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=0
    )


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = INK_900

    # Build each section
    build_cover_page(doc)
    build_toc(doc)
    build_brand_overview(doc)
    build_brand_pillars(doc)
    build_logo_section(doc)
    build_color_palette(doc)
    build_typography(doc)
    build_voice_tone(doc)
    build_photography(doc)
    build_ui_patterns(doc)
    build_social_media(doc)
    build_contact_assets(doc)

    doc.save(OUTPUT)
    print(f"Brand guide saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
